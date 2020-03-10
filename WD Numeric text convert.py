# -*- coding: utf-8 -*-
"""
Created on Mon Feb 24 14:48:07 2020

@author: andrew.white
"""

# import pprint
import os
import pandas as pd
from docx import Document
# import string
# from collections import Counter
# from nameparser import HumanName
import re
import home_finder

mydir = r"C:\Users\andrew.white\Documents\Lists\WD Numeric"
docs = os.listdir(mydir)
Sources = set()
records = []
columns = {"Source"}
rec_dict = {"Source": []}
fidx = 0
foo = []
api_key = "AIzaSyBoV_ibysiN0_yHRZq3yewXKM1bvSWwR-0"
cse = "002662449727427740006:ntnlx5npmoi"
keys_of_note = ["title", "link", "displayLink"]


def is_para_bold(file, para):
    """
    Check if entire paragrph is bold.

    It counts the bolded runs and
    compares that the to total number of runs.

    Parameters
    ----------
    para : TYPE obj
        DESCRIPTION.
        a paragraph object from the docx module
    Returns
    -------
    bold : TYPE bool
        DESCRIPTION.
    """
    bold, bcount = False, 0
    r_len = sum(1 if len(x.text) > 1 else 0 for x in para.runs)
    for r in para.runs:
        if r.bold and len(r.text) > 1:
            bcount = bcount+1
    if bcount >= r_len/2:
        bold = True
    return bold


for file in docs:
    bold_list = set()
    record_values = set()
    if file.endswith(".docx"):
        docx = Document(os.path.join(mydir, file))
        fidx = fidx + 1
        # print(file, fidx)
    else:
        continue
# record_values = set()
# file = "A.docx"
# docx = Document(os.path.join(mydir, file))
# Sources.add(file)
# fidx = fidx + 1
    pcount = 0
    for para in docx.paragraphs:
        if para.text.find(":") and len(para.text) > 1:
            Sources.add(file)
            if is_para_bold(file, para):
                pcount = pcount+1
                rec_idx = str(fidx) + "." + str(pcount)
                upd = (rec_idx, "Source", file)
                if upd not in records:
                    records.append(upd)
            c = para.text.split(": ", 1)[0]\
                .strip()\
                .rstrip(":.?,")
            try:
                v = para.text.split(": ", 1)[1].strip()
            except IndexError:
                v = "Missing delimiter in: '"+para.text+"' from file: "+file
            columns.add(c)
            record_values.add(v)
            upd = (rec_idx, c, v)
            if upd not in records:
                records.append((rec_idx, c, v))
df = pd.DataFrame(records, columns=["idx", "cols", "vals"])
df.to_csv(os.path.join(mydir, "raw.csv"))
df["cols"].replace({"Name": "Company Name"}, inplace=True)
df = df.pivot_table(index="idx", columns="cols", values="vals",
                    aggfunc=lambda x: ", ".join(str(v) for v in x))
raw_contacts = df.Contact.str\
    .replace("President & CEO|President&CEO", "President/CEO", regex=True)


def split_folk(folk, delims):
    """
    Split a sting of multiple names in a list of names, based on passed delims.

    Parameters
    ----------
    folk : STRING
        The sting containg the names.
    delimes: LIST
        The delims to split by.

    Returns
    -------
    folk_list : LIST
        The string devied by the delims.

    """
    spliter = "|".join(map(lambda x: re.escape(x), delims))
    folk_list = re.split(spliter, folk, flags=re.IGNORECASE)
    return folk_list


test = 10
spliter_list = [" or ", ","]
split_contacts = raw_contacts.apply(split_folk, args=(spliter_list,))
merge = split_contacts.apply(pd.Series)
merge = merge.rename(columns=lambda x: "Contact." + str(x))
df = pd.concat([df.drop(["Contact"], axis=1), merge[:]], axis=1)
if test > 0:
    company_names = df.sample(test)["Company Name"]
else:
    company_names = df["Company Name"]
homepages = home_finder.finder(company_names, api_key, cse, keys_of_note)
hp = pd.DataFrame.from_dict(homepages, orient="index",
                            columns=["URL", "URL Strength", "URL Ratio"])
df = df.merge(hp, left_on="Company Name", right_index=True)
df.to_csv(os.path.join(mydir, "results.csv"))


# words = [item
#          for sublist in [w.split()
#                          for w in df["Contact"]]
#          for item in sublist]
# titles = {HumanName(name).title for name in words}
# print(titles)
# punc_table = str.maketrans('', '', string.punctuation)
# bag = list([wp.translate(punc_table) for wp in words])
# bag_count = Counter(bag)
# print(bag_count.most_common())
